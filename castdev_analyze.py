#!/usr/bin/env python3
"""
Транскрибирует кастдев-интервью через Deepgram,
анализирует через Claude CLI и сохраняет отчёт в DOCX.
"""

import os
import subprocess
import sys
from pathlib import Path

import httpx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DEEPGRAM_API_KEY = "eb0f6e498007e1ee160a168f76a9f4515df23b11"
BASE_DIR = Path("/Users/vas/CLAUDECODE")
TRANS_DIR = BASE_DIR / "transcriptions"
OUTPUT_DOCX = BASE_DIR / "castdev_analysis.docx"

FILES = sorted(BASE_DIR.glob("*.m4a"))


# ─── Deepgram ─────────────────────────────────────────────────────────────────

def transcribe(path: Path) -> str:
    print(f"  🎤 Транскрибирую {path.name} ({path.stat().st_size // 1024}KB)…")
    with open(path, "rb") as f:
        audio = f.read()
    resp = httpx.post(
        "https://api.deepgram.com/v1/listen",
        params={
            "model": "nova-2",
            "language": "ru",
            "smart_format": "true",
            "punctuate": "true",
            "paragraphs": "true",
        },
        headers={
            "Authorization": f"Token {DEEPGRAM_API_KEY}",
            "Content-Type": "audio/mp4",
        },
        content=audio,
        timeout=300,
    )
    resp.raise_for_status()
    data = resp.json()
    # Пробуем достать текст с параграфами, иначе — plain transcript
    try:
        paragraphs = data["results"]["channels"][0]["alternatives"][0]["paragraphs"]["transcript"]
        return paragraphs.strip()
    except (KeyError, TypeError):
        return data["results"]["channels"][0]["alternatives"][0]["transcript"].strip()


# ─── Claude ───────────────────────────────────────────────────────────────────

def ask_claude(prompt: str) -> str:
    print("  🤖 Анализирую через Claude…")
    result = subprocess.run(
        ["claude", "-p", prompt, "--dangerously-skip-permissions"],
        capture_output=True,
        text=True,
        timeout=300,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or "claude error")
    return result.stdout.strip()


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def make_docx(analysis: str, trans_count: int):
    doc = Document()

    # Заголовок
    title = doc.add_heading("Кастдев-анализ: голос клиента", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Проанализировано интервью: {trans_count}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Парсим секции из ответа Claude (разделяем по заголовкам ## или **)
    current_heading = None
    for line in analysis.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("## ") or line.startswith("# "):
            heading_text = line.lstrip("#").strip()
            doc.add_heading(heading_text, level=1)
        elif line.startswith("**") and line.endswith("**"):
            doc.add_heading(line.strip("*").strip(), level=2)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.lstrip("-•").strip())
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x99)
        else:
            doc.add_paragraph(line)

    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Сохранено: {OUTPUT_DOCX}")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    TRANS_DIR.mkdir(exist_ok=True)

    # 1. Транскрибация
    transcriptions = {}
    for f in FILES:
        txt_path = TRANS_DIR / f"{f.stem}.txt"
        if txt_path.exists() and txt_path.stat().st_size > 100:
            print(f"  ✅ {f.name} — уже есть, пропускаю")
            transcriptions[f.stem] = txt_path.read_text(encoding="utf-8")
        else:
            try:
                text = transcribe(f)
                txt_path.write_text(text, encoding="utf-8")
                transcriptions[f.stem] = text
                print(f"     → {len(text)} символов")
            except Exception as e:
                print(f"  ❌ Ошибка {f.name}: {e}")

    if not transcriptions:
        print("Нет транскрипций — выход")
        sys.exit(1)

    print(f"\n📄 Транскрипций готово: {len(transcriptions)}")

    # 2. Анализ
    combined = "\n\n".join(
        f"=== ИНТЕРВЬЮ {k} ===\n{v}"
        for k, v in sorted(transcriptions.items())
    )

    analysis_prompt = f"""Ты аналитик кастдев-интервью. Перед тобой {len(transcriptions)} транскрипций интервью клиентов массажиста.

{combined}

Проанализируй ВСЕ интервью вместе и структурируй отчёт по блокам:

## 1. Повторяющиеся боли и запросы
Топ-5 с указанием частоты (сколько раз встречается). Повторение = сигнал.

## 2. Слова клиентов про специалиста (УТП)
Дословные цитаты — именно так они описывают мастера своими словами. Это золото для маркетинга.

## 3. Причины выбора этого массажиста
Что конкретно склонило к записи? Паттерны.

## 4. Причины возвращаться
Что удерживает клиентов? Конкретные формулировки.

## 5. Возражения и страхи до записи
Что мешало записаться раньше? Что преодолевали?

## 6. Результаты и ценность
Что изменилось после? Как клиенты описывают эффект?

## 7. Готовые цитаты для постов и шапки профиля
10-15 дословных цитат, которые звучат живо и убедительно. Без правки.

## 8. Портрет идеального клиента
Пол, возраст, ситуация, боль, триггер к записи, ожидания.

ВАЖНО: выделяй паттерны, не общие слова. Конкретика и цитаты — основа анализа."""

    try:
        analysis = ask_claude(analysis_prompt)
    except Exception as e:
        print(f"❌ Ошибка Claude: {e}")
        sys.exit(1)

    # 3. DOCX
    make_docx(analysis, len(transcriptions))


if __name__ == "__main__":
    main()
